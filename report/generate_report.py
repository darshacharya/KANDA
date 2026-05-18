"""
KANDA Project Report Generator — Full Version (30+ pages)
Run: python3 generate_report.py
Output: kanda_report.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
import os

doc = Document()

IMGS = os.path.join(os.path.dirname(__file__), "..", "imgs")
IMG_ROBOT1 = os.path.join(IMGS, "kanda(1).jpeg")
IMG_ROBOT2 = os.path.join(IMGS, "kanda(2).jpeg")
IMG_ARCH   = os.path.join(IMGS, "system-architecture.png")
IMG_LLD    = os.path.join(IMGS, "pin-level-lld.png")
IMG_FLOW   = os.path.join(IMGS, "flowchart.png")

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.5)
    section.right_margin  = Inches(1.0)

# ── Style helpers ─────────────────────────────────────────────────────────────

def chapter_label(text):
    """Chapter number line — left-justified, 16pt bold (per guidelines)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p


def heading(text, level=1):
    """level=1 → section (16pt), level=2/3 → subsection (14pt)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p


def table_caption(text):
    """Table caption — placed ABOVE the table, chapter-numbered (e.g. Table 3.1)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def para(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing_rule = 1   # ONE_POINT_FIVE
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.font.size  = Pt(12)
    run.font.name  = "Times New Roman"
    return p


def objective(text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def page_break():
    doc.add_page_break()


def center(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(8)
    return p


def figure(path, caption_text, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(caption_text)
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"
    r.font.italic = True


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1B2A4A')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # data rows
    for ri, row in enumerate(rows):
        drow = t.rows[ri + 1]
        fill = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cell = drow.cells[ci]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t


# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
center("MAJOR PROJECT REPORT", size=16)
center("on", size=13, bold=False)
doc.add_paragraph()
center("KANDA", size=22)
center("Knowledge-driven Autonomous Navigation and Decision-making Agent", size=13)
doc.add_paragraph()
center("A Multimodal Embodied Robot Agent Powered by Large Language Models", size=14)
center("with Hardware Aware Action Generation", size=14)
doc.add_paragraph()
center("Submitted in partial fulfilment of the requirements", size=12, bold=False)
center("for the award of the degree of", size=12, bold=False)
center("Bachelor of Engineering", size=13)
doc.add_paragraph()
center("Department of [Your Department]", size=12, bold=False)
center("[Your Institution Name]", size=13)
center("[Academic Year]", size=12, bold=False)
doc.add_paragraph()
doc.add_paragraph()
center("Submitted by", size=12, bold=False)
center("[Student Name(s)]", size=13)
center("[USN / Roll Number]", size=12, bold=False)
doc.add_paragraph()
center("Under the guidance of", size=12, bold=False)
center("[Guide Name], [Designation]", size=13)
center("[Department], [Institution]", size=12, bold=False)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═════════════════════════════════════════════════════════════════════════════

center("ABSTRACT", size=18)

para(
    "Household companion robotics is an emerging domain with significant social relevance, "
    "driven by the global increase in elderly populations living alone and the growing demand "
    "for accessible assistive technology. While large language model (LLM) reasoning has been "
    "integrated into robotic systems such as SayCan and RT-2, these implementations relied on "
    "expensive, proprietary hardware platforms inaccessible for broad academic or household "
    "deployment. Low-cost systems built on microcontrollers demonstrated voice-based LLM "
    "interaction but lacked autonomous navigation capability. A further unresolved issue was "
    "the absence of a validated interface between LLM-generated decisions and physical motor "
    "execution, creating a safety gap when raw model outputs were applied directly to hardware. "
    "This project, named KANDA — Knowledge-driven Autonomous Navigation and Decision-making "
    "Agent — was undertaken to address this gap by designing and implementing a two-layer "
    "robotic system combining autonomous navigation on commodity embedded hardware with a "
    "fully specified LLM reasoning architecture and a hardware-aware safety validation layer."
)

para(
    "The project followed an incremental, hardware-first development methodology across two "
    "phases. Phase 2, which was fully implemented, built the embodiment layer on an ESP32 "
    "microcontroller, integrating three HC-SR04 ultrasonic sensors with voltage divider "
    "signal conditioning, a TB6612FNG dual motor driver controlled via Pulse Width Modulation "
    "(PWM) using the Arduino core version 3 ledcAttach application programming interface (API), "
    "and an SSD1306 Organic Light-Emitting Diode (OLED) display for real-time feedback. A "
    "stable power architecture separated regulated microcontroller supply from high-current "
    "motor supply to prevent voltage instability. The intelligence layer, designed for Phase 3 "
    "on a Raspberry Pi, specified a multimodal context builder combining camera frames, "
    "OpenAI Whisper speech transcription, and sensor telemetry into a structured hardware "
    "description prompt submitted to a GPT-4 or Gemini LLM API, with a safety validator "
    "enforcing hardware-safe parameter ranges before transmitting validated JavaScript Object "
    "Notation (JSON) commands to the ESP32 over Universal Asynchronous Receiver-Transmitter "
    "(UART). The completed Phase 2 system demonstrated reliable autonomous obstacle avoidance "
    "with smooth differential steering and real-time OLED feedback, providing a tested "
    "foundation for Phase 3 LLM integration."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("CHAPTER 1")
center("INTRODUCTION", size=18)

heading("1.1 Overview of Project Topic")

para(
    "The project titled KANDA: A Multimodal Embodied Robot Agent Powered by LLMs with Hardware "
    "Aware Action Generation was carried out to investigate how large language model reasoning "
    "could be meaningfully integrated with a physically embodied robotic platform built on "
    "commodity hardware. The word KANDA is an acronym that expands to Knowledge-driven "
    "Autonomous Navigation and Decision-making Agent, a title that reflects both the "
    "technical ambition of the project and its emphasis on combining physical autonomy "
    "with AI-driven reasoning."
)

para(
    "The central insight driving the project was that household companion robotics represented "
    "a uniquely appropriate domain for LLM integration. Unlike industrial or agricultural "
    "robotic applications, where rule-based logic can often handle the full range of "
    "expected situations, household companion tasks involve open-ended human interaction, "
    "unpredictable conversational inputs, and a wide variety of contextual demands that "
    "cannot be pre-programmed. A robot that could converse with an elderly person, adapt "
    "its language for a child, remember medication schedules, respond to voice commands "
    "for home automation, and detect emergency situations required not just fast sensor "
    "processing but genuine contextual reasoning. Large language models were identified as "
    "the most appropriate technology for providing this reasoning capability."
)

para(
    "At the same time, the project recognised that the existing landscape of LLM-powered "
    "robotic systems was dominated by expensive, specialised hardware platforms that were "
    "inaccessible for widespread deployment. Research systems such as SayCan and RT-2 "
    "demonstrated the feasibility of embodied AI but required hardware investment beyond "
    "the reach of most educational, research, and household deployment contexts. KANDA was "
    "therefore positioned as a contribution to the democratisation of intelligent robotics, "
    "demonstrating that the core capabilities of an LLM-powered companion robot could be "
    "realised on commodity components costing a small fraction of research platform costs."
)

para(
    "The project followed a structured phased development plan. Phase 1 covered conceptual "
    "design, literature review, and architecture planning. Phase 2, which formed the "
    "completed portion of this work, built and validated the embodiment layer. Phase 3 "
    "defined the intelligence layer integration and was ready to be implemented in the "
    "next stage of the project. This report documents all phases in full, with emphasis "
    "on the engineering decisions, technical challenges, and solutions developed during "
    "Phase 2."
)

heading("1.1.1 Global Scenario")

para(
    "Globally, the integration of artificial intelligence with physical robotic systems "
    "has advanced considerably since the early 2020s. The publication of SayCan by Ahn "
    "et al. in 2022 established that large language models could be used to decompose "
    "high-level user instructions into robot-executable skill sequences, using the "
    "physical affordances of the robot's environment to constrain which actions were "
    "feasible. RT-2, published by Brohan et al. in 2023, extended this further by "
    "training a vision-language model that could directly output robot action commands "
    "from image and text input, effectively eliminating the need for a separate action "
    "library. PaLM-E, published by Driess et al. in 2023, demonstrated that embodied "
    "sensor data could be included directly in the input context of a large multimodal "
    "model, enabling a single unified model to handle both language understanding and "
    "physical world interaction."
)

para(
    "In parallel, open-source communities have begun exploring LLM integration on "
    "commodity hardware. ESP-SparkBot, demonstrated by Espressif in April 2025, showed "
    "that an ESP32-S3 microcontroller could be used to build a voice-interactive robot "
    "with ChatGPT integration, facial recognition, and smart home connectivity. Seervo, "
    "another open-source project, demonstrated that an ESP32 could capture images, "
    "send them to a GPT-5 API server, and receive JSON movement commands executed in "
    "MicroPython. OpenCastor, a broader framework, showed that multiple hardware "
    "platforms including the Raspberry Pi, ESP32, and Arduino could be connected to "
    "various LLM providers through a YAML configuration system. These developments "
    "confirmed that LLM integration at the commodity hardware level was technically "
    "feasible, while also revealing that navigation-capable, companion-focused systems "
    "with safety-validated command interfaces remained largely absent from the open-source "
    "landscape."
)

para(
    "In the companion robotics domain, research published in JMIR Human Factors in 2025 "
    "evaluated socially assistive robots integrated with large language models in hospital "
    "geriatric units and identified gaps in personalisation, response latency, and "
    "accessibility for low-cost deployment. A participatory design study published in "
    "Frontiers in Robotics and AI in 2024 engaged older adults in the co-design of "
    "conversational companion robots and identified key expectations including active "
    "engagement during isolation, memory of previous conversations, and emotional "
    "expression. These findings collectively indicated that companion robotics was a "
    "high-value, underserved application domain where affordable LLM-powered systems "
    "could make a significant contribution."
)

heading("1.1.2 Societal Relevance")

para(
    "The social relevance of this project was grounded in two major demographic and "
    "technological trends. The first was the global increase in the elderly population "
    "living alone or with reduced access to regular human caregiving. In many countries, "
    "the combination of longer lifespans and declining birth rates has created a growing "
    "cohort of older adults who require daily support that family members or professional "
    "caregivers cannot always provide. A companion robot capable of conversational "
    "interaction, health reminders, and emergency response could address a portion of "
    "this gap in a way that was accessible, always available, and not limited by "
    "geographic distance from family members."
)

para(
    "The second trend was the expansion of home automation and smart device ecosystems, "
    "which have made household environments increasingly amenable to robot integration. "
    "Existing voice assistants and smart home platforms could already control lights, "
    "thermostats, and appliances through natural language commands, but they lacked "
    "physical mobility and contextual embodiment. A companion robot that could navigate "
    "the home environment, identify the user's location, and interact with both the user "
    "and smart home devices represented a significant step toward a fully integrated "
    "household AI presence. KANDA was designed with this integration potential in mind, "
    "with its two-layer architecture providing the physical navigation capability and the "
    "LLM reasoning capacity required for such an integration."
)

para(
    "Beyond elderly assistance, the project also addressed the educational context through "
    "its child tutoring use case. An affordable companion robot capable of adapting its "
    "language and explanation style to a child's level, answering spontaneous questions, "
    "and maintaining contextual continuity across a tutoring session could serve as a "
    "supplement to formal education in environments where one-on-one teaching time was "
    "limited. The availability of such a system on a low-cost hardware platform made "
    "it feasible for use in homes across a wide socioeconomic range, contributing to a "
    "more equitable distribution of educational technology."
)

heading("1.1.3 Problem Area")

para(
    "Despite the advances in both large language models and robotic hardware, the "
    "specific problem of integrating LLM reasoning with a navigating, low-cost embedded "
    "robot for household companion applications remained largely unaddressed. The "
    "available LLM-powered robotic systems fell into two categories: expensive research "
    "platforms with sophisticated hardware but inaccessible cost structures, and "
    "low-cost commodity systems that provided voice interaction but no autonomous "
    "navigation capability. Neither category addressed the intersection of navigation, "
    "companion use cases, and low-cost reproducibility that KANDA targeted."
)

para(
    "A further dimension of the problem was the absence of a validated command interface "
    "between LLM output and physical motor execution. Large language models generated "
    "text or structured data responses, but they operated without any inherent knowledge "
    "of GPIO voltage limits, motor current ratings, PWM resolution, or the mechanical "
    "constraints of a specific robot platform. Applying LLM output directly to motor "
    "control without a validation layer could produce commands that were physically "
    "impossible, outside safe operational ranges, or potentially damaging to hardware. "
    "Research on safety guardrails for LLM-enabled robots, such as the RoboGuard system "
    "published in 2025, demonstrated that unconstrained LLM commands led to unsafe "
    "behaviour in up to 92 percent of worst-case scenarios. This problem was directly "
    "addressed in the KANDA design through the safety validation layer specified in the "
    "intelligence layer architecture."
)

para(
    "The problem area therefore combined three distinct challenges: building a reliable "
    "hardware execution platform on commodity components, designing an effective multimodal "
    "LLM reasoning pipeline for companion interactions, and creating a validated interface "
    "that allowed LLM-generated commands to be safely executed by the physical hardware. "
    "These three challenges together defined the scope and contribution of the KANDA project."
)

heading("1.2 Specific Details of the Project Topic")

para(
    "This project dealt with the design and implementation of an embodied robotic agent "
    "intended for indoor household companion use. The topic was specific in its "
    "combination of three technical elements that had not previously been integrated in "
    "a single open, reproducible system: autonomous navigation on commodity embedded "
    "hardware, multimodal LLM reasoning through a Raspberry Pi bridge, and a "
    "hardware-aware JSON command protocol with safety validation. The project was not a "
    "theoretical study of any one of these elements in isolation, but an engineering "
    "effort to combine all three into a functioning system."
)

para(
    "The hardware platform selected for the project reflected the constraint of "
    "reproducibility and low cost. The ESP32 microcontroller was chosen as the execution "
    "unit because of its dual-core processing architecture, 240 MHz clock speed, built-in "
    "WiFi and Bluetooth connectivity, 34 available GPIO pins, and support for PWM-based "
    "motor control through the Arduino core version 3 API. The ESP32 DevKit form factor "
    "was used because of its convenient USB programming interface and the availability of "
    "3.3-volt regulated output for sensor interfacing. The total cost of the complete "
    "Phase 2 hardware assembly was significantly lower than any comparable research "
    "robot platform."
)

para(
    "The intelligence layer was designed around the Raspberry Pi 4 Model B or equivalent, "
    "chosen because it ran a full Linux operating system, supported Python 3.9 or later, "
    "had a dedicated camera serial interface, USB ports for microphone connection, and "
    "hardware UART pins for communication with the ESP32. The planned LLM integration "
    "used the OpenAI GPT-4 API or the Google Gemini API, both of which supported "
    "multimodal inputs including text and image data. OpenAI Whisper was selected for "
    "speech-to-text conversion because of its strong performance on diverse audio "
    "conditions without per-request API costs."
)

para(
    "A key technical detail of the project was the design of the hardware description "
    "prompt. This was a structured natural language string included in every LLM API "
    "request that described the robot's physical capabilities in terms the model could "
    "reason about. The prompt specified the available movement commands, the valid speed "
    "range, the sensor layout, and the current sensor readings. This approach was "
    "inspired by the affordance-based reasoning in SayCan but implemented without a "
    "separate skill library, relying instead on the LLM's inherent reasoning capability "
    "to generate appropriate JSON commands within the described hardware constraints."
)

heading("1.3 Problem Statement")

para(
    "Affordable and reproducible intelligent companion robots for household use do not "
    "currently exist as open, well-documented, navigation-capable systems. While "
    "large language model-powered robots have been demonstrated in research, they rely "
    "on expensive, proprietary hardware platforms inaccessible for broad deployment. "
    "Low-cost platforms built on microcontrollers such as the ESP32 have demonstrated "
    "LLM integration for voice interaction, but these systems lack autonomous navigation "
    "and are not designed for companion use cases involving elderly users or children. "
    "Furthermore, no validated open protocol exists for translating LLM-generated "
    "decisions into hardware-safe motor commands for embedded robotic systems, "
    "creating a significant safety and reliability gap in the field."
)

para(
    "The problem addressed in this project was therefore the absence of an integrated, "
    "low-cost robotic system that combined autonomous navigation, multimodal LLM "
    "reasoning, and validated hardware-aware action generation for household companion "
    "applications. KANDA was designed to address this problem by building a stable "
    "embodiment layer on commodity components, demonstrating reliable autonomous "
    "navigation, and defining a complete architecture for attaching LLM reasoning to "
    "the physical platform through a safety-validated command protocol that constrained "
    "all LLM output to hardware-safe operational parameters before execution."
)

heading("1.4 Objectives")

para(
    "The objectives of the KANDA project were defined to guide development across both "
    "phases and to ensure that each deliverable contributed to the overall goal of "
    "a hardware-grounded, LLM-powered companion robotic system."
)

objective(
    "To design and implement a stable hardware platform using the ESP32 microcontroller, "
    "TB6612FNG motor driver, three HC-SR04 ultrasonic sensors with voltage divider "
    "signal conditioning, and SSD1306 OLED display, capable of autonomous obstacle "
    "avoidance and real-time sensor feedback."
)
objective(
    "To design a hardware description prompt that communicates the robot's available "
    "sensors and actuators to a large language model as structured natural language "
    "context, enabling the LLM to reason about physically feasible actions."
)
objective(
    "To build a speech transcription pipeline using OpenAI Whisper that converts "
    "microphone input to text for submission to the large language model as part of "
    "the multimodal context."
)
objective(
    "To integrate a camera capture module that encodes visual frames and includes them "
    "in the multimodal LLM API request for scene-aware reasoning about the robot's "
    "immediate environment."
)
objective(
    "To implement a motor actuation layer on the Raspberry Pi that executes movement "
    "commands derived from the parsed LLM JSON response, following validation by a "
    "hardware-aware safety checking layer."
)
objective(
    "To integrate a text-to-speech module that enables the robot to produce verbal "
    "responses returned by the large language model, supporting conversational "
    "companion interaction with elderly users and children."
)
objective(
    "To evaluate the integrated system in household companion scenarios including elderly "
    "assistance, child tutoring, schedule reminders, home automation control, and "
    "emergency alerting, measuring response latency, command safety, and interaction quality."
)

heading("1.5 Scope of the Project")

para(
    "The scope of the present work covered two distinct phases of development. Phase 2, "
    "which was fully completed, included the hardware platform design, power architecture, "
    "component selection, sensor integration with voltage divider signal conditioning, "
    "conflict-free GPIO pin mapping, PWM-based motor control firmware, obstacle avoidance "
    "logic, and OLED display feedback. This phase produced a fully functional, standalone "
    "autonomous robot capable of navigating an indoor environment without human "
    "intervention and without any AI components, providing a stable and tested foundation "
    "for Phase 3 integration."
)

para(
    "The scope also included the complete architectural design of Phase 3, covering the "
    "Raspberry Pi UART bridge, the multimodal context builder, the hardware description "
    "prompt design, the LLM API client, the JSON command parser, the safety validator, "
    "and the text-to-speech module. Although Phase 3 components were not implemented "
    "at the time of this report, their design was sufficiently detailed to constitute "
    "a meaningful architectural contribution and to provide an unambiguous implementation "
    "roadmap. The firmware for the ESP32 was also designed with Phase 3 in mind, with "
    "GPIO pins reserved for future UART communication with the Raspberry Pi."
)

para(
    "The scope explicitly excluded large-scale user studies, clinical evaluation of "
    "elderly assistance capabilities, deployment in multiple household types, and "
    "outdoor or non-domestic environments. The project was also not intended to produce "
    "a commercial product or to claim that the system could replace professional "
    "caregiving. It was designed as an academic proof-of-concept system demonstrating "
    "the feasibility and architecture of a low-cost, LLM-powered companion robot with "
    "hardware-aware action generation."
)

heading("1.6 Methodology")

para(
    "The methodology of the KANDA project was grounded in a hardware-first, incremental "
    "integration philosophy. The decision to validate the physical execution layer before "
    "introducing any AI components was deliberate and justified by the experience of "
    "related projects in which hardware and software failures occurring simultaneously "
    "created debugging complexity that was difficult to resolve. By ensuring that each "
    "hardware component was verified in isolation before integration with the next, "
    "the project maintained a clear chain of accountability for any issue that arose."
)

para("Step 1: Literature Review and Domain Selection")
para(
    "A systematic review of related work in embodied AI, LLM-robot integration, "
    "companion robotics, and commodity hardware implementations was conducted. Key "
    "systems including SayCan, RT-2, PaLM-E, ESP-SparkBot, Seervo, RoboGuard, and "
    "companion robot studies from JMIR and Frontiers were examined. The original "
    "greenhouse monitoring domain was found to be a poor fit for LLM integration, "
    "since agricultural tasks could be adequately addressed by rule-based logic. "
    "The domain was revised to household companion robotics, where the need for "
    "LLM reasoning was strongly justified by the unpredictability and diversity of "
    "human companion interactions."
)

para("Step 2: Architecture Design")
para(
    "The two-layer architecture was designed, separating the embodiment layer on the "
    "ESP32 from the intelligence layer on the Raspberry Pi. The communication protocol "
    "between the two layers was specified as a UART link transmitting JSON-formatted "
    "commands from Pi to ESP32 and sensor telemetry from ESP32 to Pi. The hardware "
    "description prompt format was designed to provide the LLM with the robot's "
    "physical constraints in a structured, predictable way."
)

para("Step 3: Hardware Assembly and Power Architecture")
para(
    "The power system was assembled first: lithium polymer battery to battery management "
    "system to main switch, with one branch to the buck converter for regulated five-volt "
    "ESP32 power, and a second direct branch to the TB6612FNG motor driver for high-current "
    "motor supply. All grounds were tied together to form a common reference. Capacitors "
    "were placed across motor supply lines to suppress voltage spikes during motor switching."
)

para("Step 4: Sensor Integration and Signal Conditioning")
para(
    "Three HC-SR04 sensors were connected with voltage dividers on each ECHO pin. The "
    "divider ratio of one-kiloohm plus one-kiloohm was used to reduce the five-volt "
    "sensor output to approximately two-point-five volts, safely within the three-point-"
    "three volt GPIO input range of the ESP32. Input-only GPIO pins 34, 35, and 32 were "
    "selected for the ECHO signals to avoid any conflict with output-capable GPIO "
    "functions. Each sensor was tested individually before proceeding."
)

para("Step 5: Firmware Development and Integration")
para(
    "The ESP32 firmware was written in C++ using the Arduino framework and Arduino core "
    "version 3. GPIO pin assignments were validated against the ESP32 datasheet to avoid "
    "strapping pins, internal flash pins, and UART communication pins. The OLED display "
    "was initialised with an explicit Wire.begin call specifying GPIO 21 and GPIO 22. "
    "PWM was configured using the new ledcAttach API. The complete obstacle avoidance "
    "loop was assembled after each subsystem was independently verified."
)

para("Step 6: Intelligence Layer Design")
para(
    "The complete Phase 3 pipeline was designed and documented in sufficient detail for "
    "immediate implementation in the next project phase. This included the UART bridge "
    "firmware additions, the Python multimodal context builder, the LLM API integration, "
    "the safety validator specification, and the text-to-speech integration plan."
)

heading("1.7 Organization of Report")

para(
    "Chapter 2 presents the theoretical and conceptual background of the project, covering "
    "embodied AI, large language models in robotics, hardware-aware action generation, "
    "companion robotics, and relevant related work. Chapter 3 provides the software "
    "requirement specification, describing the complete system product including functional "
    "requirements, performance specifications, software and hardware requirements, and "
    "design constraints. Chapter 4 presents the high-level design of the system, including "
    "design philosophy, architectural strategies, and the overall system architecture. "
    "Chapter 5 covers the low-level design, including detailed pin mapping, signal "
    "conditioning, firmware architecture, and the LLM command protocol specification. "
    "Chapter 6 describes the Phase 2 implementation and testing in detail, including "
    "hardware assembly, key engineering decisions, and resolved challenges. Chapter 7 "
    "presents the conclusions drawn from the completed work and the future work planned "
    "for Phase 3 and beyond."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("CHAPTER 2")
center("THEORY AND CONCEPT OF PROJECT", size=18)

heading("2.1 Introduction to Embodied AI")

para(
    "Embodied artificial intelligence is a paradigm in AI research that holds that "
    "genuine intelligence cannot be separated from physical experience. The foundational "
    "argument is that a cognitive system which only processes abstract symbols or "
    "textual inputs operates in a fundamentally different and more limited way than "
    "a system that must perceive, act upon, and respond to a real physical environment. "
    "This distinction between embodied and disembodied intelligence has become "
    "particularly relevant as large language models have demonstrated impressive "
    "reasoning capabilities in text-only domains while remaining unable to act "
    "in the physical world without additional robotic integration."
)

para(
    "In the context of robotics, embodied AI is realised through the sense-decide-act "
    "cycle. The robot continuously reads from its sensors, constructs a representation "
    "of its situation, decides on an appropriate action, and executes that action "
    "through its actuators. The physical result of the action modifies the environment, "
    "which in turn generates new sensory input, closing the loop. This cycle is "
    "fundamentally different from a chatbot interaction, where the input and output "
    "are both textual and the system has no physical consequences for its responses. "
    "An embodied agent is accountable to the laws of physics in a way that a purely "
    "language-based agent is not."
)

para(
    "KANDA was designed around this embodied AI principle. The ESP32 continuously "
    "executed the sense-decide-act-display loop, reading ultrasonic sensor distances, "
    "applying movement logic, commanding the motor driver, and updating the OLED. "
    "In Phase 2, the decision step was rule-based. In Phase 3, this step was to be "
    "replaced by LLM reasoning, transforming the system from a reactive machine into "
    "a deliberative agent capable of understanding context, following natural language "
    "instructions, and generating responses that were grounded in the robot's actual "
    "physical situation."
)

heading("2.2 Large Language Models in Robotics")

para(
    "Large language models are neural network systems trained on massive text corpora "
    "that learn to generate coherent, contextually appropriate natural language "
    "continuations of any given input. The introduction of transformer architectures "
    "in 2017 and their scaling to billions of parameters enabled capabilities such as "
    "few-shot reasoning, chain-of-thought problem solving, code generation, and "
    "multimodal understanding that had not been achievable with earlier natural language "
    "processing approaches. These capabilities made LLMs attractive as reasoning "
    "components in robotic systems, where the robot's physical context could be "
    "described in natural language and the LLM could reason about appropriate actions."
)

para(
    "The SayCan system, published by Ahn et al. in 2022, was a landmark demonstration "
    "of LLM-guided robot behaviour. SayCan used an LLM to score candidate robot skills "
    "by their semantic relevance to a user instruction, then multiplied these scores "
    "by affordance values representing the robot's physical ability to execute each "
    "skill in its current environment. This grounding mechanism prevented the LLM from "
    "suggesting actions that were semantically appropriate but physically infeasible. "
    "KANDA adopted a related principle through its hardware description prompt, which "
    "communicated the robot's physical constraints directly to the LLM before any "
    "command was generated."
)

para(
    "RT-2, published by Brohan et al. in 2023, advanced the field further by training "
    "a single vision-language model on both internet-scale image-text data and robot "
    "demonstration data, enabling it to generalise web-acquired knowledge to novel "
    "robotic manipulation tasks. While RT-2 operated on a much larger scale than "
    "KANDA and required specialised robot hardware, it established the principle that "
    "visual grounding was essential for robot intelligence in unstructured environments, "
    "motivating the inclusion of a camera module in the KANDA Phase 3 design."
)

para(
    "More recent work on safety guardrails for LLM-enabled robots, including RoboGuard "
    "published in 2025 and SafeGate proposed in 2026, demonstrated that raw LLM outputs "
    "could not be trusted for direct hardware execution. RoboGuard showed that without "
    "safety constraints, up to 92 percent of adversarially crafted LLM prompts could "
    "lead to unsafe robot behaviour. These findings directly motivated the inclusion of "
    "a dedicated safety validation layer in the KANDA architecture, positioned between "
    "the LLM API response and the UART transmission to the ESP32."
)

heading("2.3 Hardware-Aware Action Generation")

para(
    "Hardware-aware action generation refers to the design of an AI-to-robot pipeline "
    "in which the output of the reasoning system is explicitly constrained by the "
    "physical capabilities and safety parameters of the hardware it controls. In a "
    "standard LLM API interaction, the model produces text based solely on the input "
    "prompt, with no knowledge of the physical system that will act on the response. "
    "For robotic applications, this creates a fundamental mismatch: the LLM may "
    "generate semantically valid commands that are physically impossible, out of range, "
    "or dangerous for the specific hardware configuration."
)

para(
    "KANDA addressed this problem through two complementary mechanisms. The first was "
    "the hardware description prompt, which informed the LLM of the available commands, "
    "valid parameter ranges, sensor layout, and current sensor readings before asking "
    "it to generate an action. By making the hardware constraints part of the LLM's "
    "input context, the probability of receiving physically reasonable commands was "
    "increased. This approach was analogous to providing an engineer with a system "
    "specification before asking for design decisions, rather than asking for decisions "
    "in the absence of any specification."
)

para(
    "The second mechanism was the safety validator, implemented on the Raspberry Pi "
    "as a post-processing step applied to every JSON command before UART transmission. "
    "The validator checked that the action field contained a recognised value from the "
    "defined command set and that the speed field, if present, fell within the "
    "zero-to-two-hundred-and-fifty-five range of the eight-bit PWM resolution. Commands "
    "failing validation were discarded and the robot was commanded to stop safely. "
    "Together, the hardware description prompt and the safety validator formed the "
    "hardware-aware action generation pipeline that distinguished KANDA from simpler "
    "LLM-to-motor implementations."
)

heading("2.4 Companion Robotics for Household Applications")

para(
    "Companion robotics is a field within human-robot interaction research concerned "
    "with robots that provide social, communicative, and assistive functions in domestic "
    "or healthcare settings. Early companion robots such as PARO, a therapeutic seal "
    "robot used in dementia care, demonstrated that robotic companions could provide "
    "measurable psychological benefits to elderly users even without sophisticated AI "
    "reasoning capabilities. The introduction of large language models to companion "
    "robotics has opened the possibility of more natural, adaptive, and contextually "
    "aware interactions that earlier systems could not achieve."
)

para(
    "Research published in JMIR Human Factors in 2025 evaluated a socially assistive "
    "robot integrated with a large language model in a hospital geriatric unit and "
    "found positive user perceptions of ease of use, usefulness, and enjoyment. "
    "However, the study identified ongoing challenges including response latency, "
    "voice quality, and the difficulty of personalising interactions for users with "
    "varying cognitive abilities. A participatory co-design study published in "
    "Frontiers in Robotics and AI in 2024 engaged 28 older adults in the design "
    "of a conversational companion robot and identified the desire for active "
    "engagement, memory of previous conversations, privacy protection, and emotional "
    "expression as key requirements. These findings directly informed the KANDA "
    "companion use case design and evaluation criteria."
)

para(
    "The household companion use cases selected for KANDA were chosen to cover a "
    "range of interaction types that required different levels of LLM reasoning "
    "sophistication. Medication reminders required temporal awareness and "
    "personalisation but relatively simple conversational interaction. Child tutoring "
    "required adaptive language, unpredictable question handling, and contextual "
    "continuity across multiple exchanges. Elderly companionship required empathetic "
    "conversational capability and recognition of emotional states. Home automation "
    "required natural language command interpretation and integration with external "
    "smart home APIs. Emergency alerting required real-time environmental assessment "
    "and decision-making under uncertainty. Together, these use cases provided a "
    "comprehensive evaluation framework for the Phase 3 intelligence layer."
)

heading("2.5 Embedded Systems for Robotic Control")

para(
    "The ESP32 microcontroller used in KANDA belongs to the class of embedded systems "
    "designed for real-time control applications. Unlike general-purpose computers, "
    "embedded systems are optimised for deterministic timing, low power consumption, "
    "and direct hardware interfacing. The ESP32's dual-core Xtensa LX6 processor "
    "architecture, operating at up to 240 MHz, provided sufficient computational "
    "throughput for the sensor reading, PWM generation, OLED rendering, and serial "
    "communication tasks required by the embodiment layer, while its large selection "
    "of GPIO pins with multiple alternative functions made it suitable for the "
    "diverse peripheral requirements of the robot."
)

para(
    "The TB6612FNG motor driver used in KANDA was designed specifically for small "
    "brushed DC motor control. Its dual H-bridge architecture allowed independent "
    "speed and direction control for two motors, which was essential for differential "
    "steering. The driver accepted logic-level inputs from the ESP32's 3.3-volt GPIO "
    "pins without requiring additional level translation, simplifying the wiring. "
    "Its built-in standby mode, activated by pulling the STBY pin high, provided "
    "a clean enable/disable mechanism, and its internal protection circuitry guarded "
    "against thermal overload and short circuits."
)

heading("2.6 Summary")

para(
    "This chapter presented the conceptual and technical foundations of the KANDA "
    "project across five areas. Embodied AI established the theoretical justification "
    "for physical grounding of intelligence. Large language models in robotics provided "
    "the context for the planned Phase 3 integration and identified the specific "
    "prior systems that KANDA built upon and differentiated from. Hardware-aware "
    "action generation explained the design principle behind the hardware description "
    "prompt and safety validation layer. Companion robotics for household applications "
    "justified the application domain and the specific use cases chosen for evaluation. "
    "Embedded systems for robotic control provided the technical background for the "
    "hardware components selected for Phase 2. Together, these concepts formed the "
    "complete theoretical foundation for the system designed and implemented in the "
    "following chapters."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("CHAPTER 3")
center("SOFTWARE REQUIREMENT SPECIFICATION FOR KANDA", size=18)

heading("3.1 Overall Description")

para(
    "The proposed system, named KANDA, was developed as a two-layer embedded robotic "
    "platform with a completed hardware execution layer and a fully specified intelligence "
    "layer awaiting implementation. The Phase 2 product was a standalone autonomous "
    "obstacle-avoidance robot operating entirely on the ESP32 microcontroller with a "
    "deterministic firmware loop. The complete product, as defined in this specification, "
    "was a household companion robotic agent in which a Raspberry Pi handled multimodal "
    "LLM reasoning and transmitted validated JSON commands to the ESP32 for hardware "
    "execution through a UART serial interface."
)

para(
    "The ESP32 firmware was written in C++ using the Arduino framework. The intelligence "
    "layer was specified to be implemented in Python on the Raspberry Pi, using the "
    "OpenAI Whisper library for speech transcription, OpenCV or the Raspberry Pi camera "
    "module library for visual capture, the OpenAI Python SDK or Google Generative AI "
    "SDK for LLM communication, pyttsx3 or gTTS for text-to-speech output, and PySerial "
    "for UART communication with the ESP32. The system operated as a closed loop in "
    "which sensor and camera data flowed from the ESP32 and Raspberry Pi peripherals "
    "to the LLM, and validated commands flowed back to the motor driver."
)

heading("3.1.1 Product Perspective")

para(
    "KANDA was designed as a self-contained embedded product, not a web service or "
    "simulation. The physical robot constituted the product itself, and all intelligence "
    "components were either co-located on the Raspberry Pi mounted on the robot or "
    "accessed through cloud API calls made by the Pi. This design ensured that the "
    "robot could operate as a standalone system in any home environment with WiFi "
    "connectivity, without requiring a separate computing device to act as a server "
    "or controller."
)

para(
    "The product perspective was structured around the hardware safety principle: the "
    "ESP32 embodiment layer could function independently of the Raspberry Pi at all "
    "times. If the Pi lost power, crashed, or lost internet connectivity, the ESP32 "
    "defaulted to its local rule-based obstacle avoidance mode, preventing the robot "
    "from becoming stationary or unsafe. This fallback behaviour was a product "
    "requirement rather than an optional feature, ensuring that the companion robot "
    "remained safe and mobile under all failure conditions."
)

heading("3.1.2 Product Functions")

para(
    "The primary product function was autonomous navigation in an indoor household "
    "environment. The robot continuously sensed distances on three axes, made movement "
    "decisions, and controlled motor speed and direction to avoid obstacles and "
    "navigate corridors, rooms, and doorways. The second function was natural language "
    "interaction, where the robot accepted spoken user input through a microphone, "
    "transcribed it using Whisper, incorporated it into the LLM context along with "
    "visual and sensor data, and generated both a verbal response through text-to-speech "
    "and a movement command."
)

para(
    "The third function was multimodal context assembly, where camera frames, "
    "microphone audio, and ESP32 sensor telemetry were combined into a structured "
    "LLM input containing the hardware description prompt, the current sensor readings, "
    "a description of the visual scene, and the transcribed user speech. The fourth "
    "function was safety-validated command execution, where every JSON response from "
    "the LLM was verified before being transmitted to the ESP32. The fifth function "
    "was real-time display feedback, where the ESP32 OLED continuously showed current "
    "sensor distances and the active action label."
)

heading("3.1.3 Constraints")

para(
    "The system was constrained to indoor household environments with flat or "
    "near-flat floor surfaces. The HC-SR04 sensors had reliable range only between "
    "two and four hundred centimetres, limiting the robot's ability to detect "
    "very close obstacles or objects at extreme distances. The two-wheel differential "
    "drive chassis provided no lateral movement capability, constraining navigation "
    "to forward, backward, and pivot-turn movements. LLM-based reasoning required "
    "internet connectivity for cloud API access, though local model deployment on "
    "the Raspberry Pi was identified as a future enhancement to remove this constraint."
)

para(
    "The safety validator was limited to numerical range checking and did not perform "
    "semantic evaluation of whether a command was contextually appropriate. This "
    "meant that a LLM command that was within the valid parameter range but contextually "
    "inappropriate, such as driving toward a user at maximum speed in a confined space, "
    "would not be blocked by the validator alone. Contextual safety was therefore "
    "dependent on the quality of the hardware description prompt and the LLM's "
    "reasoning capability, rather than on the validator."
)

heading("3.2 Specific Requirements")

heading("3.2.1 Functional Requirements", level=3)

para(
    "The system shall read distance measurements from three HC-SR04 ultrasonic sensors "
    "at a minimum update rate of ten readings per second per sensor. The system shall "
    "drive motor speed using PWM signals at one kilohertz frequency with eight-bit "
    "resolution, providing a speed range of zero to two hundred and fifty-five. The "
    "system shall display current sensor readings and the active movement decision "
    "on the OLED display, refreshing at least ten times per second. The Raspberry Pi "
    "shall transmit a validated JSON command to the ESP32 within the response latency "
    "requirements specified in section 3.2.2. The ESP32 shall parse and execute the "
    "received JSON command within fifty milliseconds of receipt."
)

para(
    "The system shall accept spoken user input through a microphone connected to the "
    "Raspberry Pi and transcribe it to text using the Whisper speech recognition model. "
    "The system shall capture camera frames from the Raspberry Pi camera module and "
    "encode them for inclusion in the LLM API request. The LLM API shall receive a "
    "structured multimodal input containing the hardware description prompt, current "
    "sensor readings, encoded camera frame, and transcribed user speech. The LLM "
    "response shall be parsed to extract a JSON command object containing an action "
    "field and optional speed field. The safety validator shall verify the action and "
    "speed values before any command is transmitted to the ESP32."
)

heading("3.2.2 Performance Requirements", level=3)

para(
    "The obstacle avoidance loop on the ESP32 shall complete a full sense-decide-act-"
    "display cycle within one hundred milliseconds under all operating conditions. "
    "The complete LLM reasoning cycle, from multimodal context assembly to validated "
    "JSON command receipt, shall complete within five seconds under standard WiFi "
    "network and cloud API conditions. The safety validator shall complete command "
    "verification within ten milliseconds. The text-to-speech module shall begin "
    "audio playback within two seconds of receiving the LLM text response. The "
    "Whisper transcription shall complete within three seconds for utterances up "
    "to ten seconds in duration."
)

heading("3.2.3 Software Requirements", level=3)

table_caption("Table 3.1 — Software and library requirements for the KANDA system")
add_table(
    ["Component", "Software / Library", "Version"],
    [
        ["ESP32 Firmware", "Arduino IDE + ESP32 Arduino Core", "v3.0+"],
        ["OLED Display", "Adafruit SSD1306 + Adafruit GFX", "Latest"],
        ["Pi OS", "Raspberry Pi OS (64-bit)", "Bookworm+"],
        ["Language (Pi)", "Python", "3.9+"],
        ["Serial Comms", "PySerial", "Latest"],
        ["Speech-to-Text", "OpenAI Whisper", "Latest"],
        ["LLM Client", "OpenAI Python SDK / Google Gen AI SDK", "Latest"],
        ["Image Capture", "OpenCV / Picamera2", "Latest"],
        ["Text-to-Speech", "pyttsx3 / gTTS", "Latest"],
    ],
    col_widths=[2.0, 3.2, 1.5]
)

heading("3.2.4 Hardware Requirements", level=3)

table_caption("Table 3.2 — Hardware component specifications for the KANDA system")
add_table(
    ["Component", "Specification", "Purpose"],
    [
        ["ESP32 DevKit", "Dual-core 240MHz, 3.3V, 34 GPIO", "Execution microcontroller"],
        ["TB6612FNG", "Dual H-bridge, 1.2A/ch, 15V max", "Motor driver"],
        ["HC-SR04 ×3", "2–400cm, 5V, 15° beam angle", "Distance sensing"],
        ["SSD1306 OLED", "128×64px, I2C, 3.3–5V", "Real-time display"],
        ["LiPo Battery", "7.4V 2S, ≥2000mAh", "Power source"],
        ["BMS", "2S Li-ion protection board", "Battery protection"],
        ["Buck Converter", "7.4V → 5V, ≥1A", "ESP32 power regulation"],
        ["Raspberry Pi", "Pi 4 Model B, ≥2GB RAM", "AI compute bridge"],
        ["Camera Module", "Raspberry Pi Camera v2 or HQ", "Visual input"],
        ["Microphone", "USB or I2S microphone module", "Audio input"],
        ["Speaker", "3.5mm or USB speaker", "TTS audio output"],
    ],
    col_widths=[2.2, 2.8, 2.4]
)

heading("3.2.5 Design Constraints", level=3)

para(
    "The primary design constraint was the protection of ESP32 GPIO input pins from "
    "the five-volt ECHO signals produced by the HC-SR04 sensors. The ESP32 GPIO "
    "pins had a maximum input voltage of 3.6 volts, and applying five volts directly "
    "would risk permanent damage to the microcontroller. This was addressed through "
    "a voltage divider on each ECHO pin, with component values selected to produce "
    "an output of approximately 2.5 volts from a 5-volt input."
)

para(
    "A further constraint was the avoidance of ESP32 GPIO pins reserved for "
    "internal functions. Pins 0, 2, 12, and 15 were strapping pins whose state "
    "at boot time affected the ESP32's flash voltage configuration and boot mode "
    "selection. Pins 1 and 3 were the UART0 TX and RX pins used by the Arduino "
    "serial monitor during development and reserved for the Raspberry Pi UART "
    "bridge in Phase 3. Pins 6 through 11 were connected to the internal SPI flash "
    "memory and could not be used for any external peripheral function. All of "
    "these pins were excluded from the Phase 2 pin assignment."
)

heading("3.3 Summary")

para(
    "The KANDA system was specified as a two-layer embedded robotic agent combining "
    "a hardware execution layer on the ESP32 and a planned intelligence layer on the "
    "Raspberry Pi. The product functions covered autonomous navigation, natural "
    "language interaction, multimodal context assembly, safety-validated command "
    "execution, and real-time display feedback. The specific requirements established "
    "precise performance targets for each system component, defined the software and "
    "hardware stack, and identified the key design constraints governing the "
    "implementation. The constraint analysis in particular demonstrated the engineering "
    "depth of the Phase 2 hardware design, where careful pin selection and signal "
    "conditioning decisions were required to build a reliable and safe embedded system."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("CHAPTER 4")
center("HIGH LEVEL DESIGN OF THE KANDA SYSTEM", size=18)

heading("4.1 Design Considerations")

heading("4.1.1 General Considerations", level=3)

para(
    "Several fundamental considerations shaped the high-level design of the KANDA "
    "system. The most important was the principle of hardware independence. The "
    "robot had to be capable of safe, useful autonomous operation even when the "
    "Raspberry Pi or the LLM API was unavailable. Designing the ESP32 firmware as "
    "a complete, standalone obstacle-avoidance system ensured that the robot would "
    "never become immobile or potentially dangerous because of a software failure in "
    "the intelligence layer. This also made the development process safer, because "
    "Phase 3 integration could be tested incrementally without risk of the robot "
    "losing basic navigational capability during development."
)

para(
    "The second consideration was physical safety through command validation. Since "
    "the LLM reasoned about actions in natural language terms without direct knowledge "
    "of the robot's hardware specifications, there was no guarantee that its "
    "JSON outputs would always fall within the safe operational parameters of the "
    "motor system. The design response was to create a dedicated safety validation "
    "layer on the Raspberry Pi that verified every command before transmission. "
    "This layer was positioned as close as possible to the physical execution point "
    "to minimise the risk of unsafe commands reaching the hardware."
)

para(
    "The third consideration was cost and reproducibility. The complete bill of "
    "materials for the Phase 2 hardware was designed to be available from standard "
    "electronics distributors, with no custom components or proprietary modules. "
    "The ESP32 and Raspberry Pi were both widely available at low cost and had "
    "large communities of developers. The choice of the Arduino framework for ESP32 "
    "firmware and Python for the Raspberry Pi intelligence layer reflected this "
    "emphasis on widely understood, well-documented technologies that other developers "
    "could readily adopt and extend."
)

heading("4.1.2 Development Methods", level=3)

para(
    "The development of KANDA followed an incremental, hardware-first methodology. "
    "Each hardware subsystem was assembled, verified, and tested before the next was "
    "added. The sequence was: power architecture, ESP32 boot verification, ultrasonic "
    "sensor integration, OLED display integration, motor driver integration, and "
    "finally the complete obstacle avoidance loop. This sequence was chosen so that "
    "each addition had only one new variable, making any failure attributable to the "
    "most recently added component."
)

para(
    "For the firmware, an iterative development approach was used in which a minimal "
    "working sketch was written for each subsystem, tested via serial monitor output, "
    "and then merged into the main firmware. This avoided the common debugging problem "
    "of introducing multiple features simultaneously and then being unable to identify "
    "which feature caused a regression. The final firmware was the result of several "
    "iterations of pin assignment refinement, API migration, and logic tuning based "
    "on observed robot behaviour."
)

heading("4.2 Architecture Strategies")

heading("4.2.1 Technology Stack", level=3)

para(
    "The ESP32 firmware was written in C++ using the Arduino framework because this "
    "combination provided a large library ecosystem, a familiar development environment "
    "for embedded developers, and strong community support for the ESP32 platform "
    "specifically. The Arduino core version 3 for ESP32 was used, which introduced "
    "the single-call ledcAttach API for PWM configuration, replacing the two-step "
    "ledcSetup and ledcAttachPin sequence from earlier versions. This API was chosen "
    "to ensure forward compatibility with future Arduino core updates."
)

para(
    "Python was selected for the Raspberry Pi intelligence layer because it provided "
    "the widest range of library support for the required components: speech "
    "recognition, image capture, LLM API access, text-to-speech, and serial "
    "communication. Python's high-level data handling also simplified the JSON "
    "parsing and validation logic required for the command interface. The "
    "single-language approach on the Pi reduced integration complexity compared "
    "to a multi-language architecture."
)

heading("4.2.2 Future Plans", level=3)

para(
    "The KANDA architecture was designed with several future extension paths in mind. "
    "The first was the replacement of cloud-based LLM API calls with a locally "
    "deployed quantised language model running on the Raspberry Pi, eliminating the "
    "internet connectivity requirement and reducing response latency. Models such "
    "as quantised versions of Llama or Mistral had been demonstrated to run on "
    "Raspberry Pi 4 hardware with acceptable performance for short conversational "
    "exchanges, and Raspberry Pi 5 improved on this significantly."
)

para(
    "The second extension path was the enhancement of the safety validator to "
    "include semantic constraint checking using a lightweight secondary model. "
    "A small classifier could evaluate whether the LLM's intended action was "
    "contextually appropriate given the sensor readings, providing an additional "
    "layer of safety beyond numerical range validation. The third path was the "
    "addition of multi-turn conversation memory, enabling the robot to maintain "
    "context across multiple interactions with the same user and provide more "
    "personalised companion behaviour. The fourth was integration with smart home "
    "APIs such as Home Assistant or Google Home, expanding the robot's ability "
    "to control household devices as part of its companion function."
)

heading("4.3 System Architecture")

para(
    "The KANDA system was designed using a two-layer embedded architecture with a "
    "bidirectional UART communication bridge between the layers. The first layer, "
    "the embodiment layer centred on the ESP32, handled all real-time hardware "
    "interaction including sensor reading, motor control, and OLED feedback. The "
    "second layer, the intelligence layer centred on the Raspberry Pi, handled all "
    "AI pipeline operations including multimodal context assembly, LLM API "
    "communication, response validation, and verbal output generation."
)

para(
    "The power architecture was designed as an independent subsystem. The lithium "
    "polymer battery supplied seven-point-four volts nominal, protected by a battery "
    "management system. A main switch controlled power distribution to two branches: "
    "a buck converter branch providing regulated five volts to the ESP32 VIN pin, "
    "and a direct branch providing full battery voltage to the TB6612FNG motor driver "
    "VM pin. This separation prevented the voltage spikes generated by motor switching "
    "from propagating back through the power supply to the microcontroller."
)

para(
    "Figure 4.1 presents the complete high-level system architecture diagram. "
    "Completed Phase 2 components are shown with green colouring, and planned "
    "Phase 3 components are shown with grey colouring, clearly communicating "
    "the implementation status of each element of the system."
)

figure(IMG_ARCH,
       "Figure 4.1 — KANDA High-Level System Architecture (green = completed, grey = planned)",
       width=Inches(5.8))

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("CHAPTER 5")
center("LOW LEVEL DESIGN OF THE KANDA EMBODIMENT LAYER", size=18)

heading("5.1 Pin Mapping and GPIO Assignment")

para(
    "The low-level design of the KANDA embodiment layer required a careful analysis "
    "of the ESP32 DevKit GPIO capabilities to produce a conflict-free pin assignment. "
    "The ESP32 contained several categories of GPIO pins that had to be handled "
    "differently: input-only pins that could not be configured as outputs, strapping "
    "pins whose state at boot time affected the chip's configuration, UART pins "
    "required for serial communication, and internal flash pins that were permanently "
    "reserved. The final pin assignment was the result of iterative refinement and "
    "datasheet review."
)

table_caption("Table 5.1 — ESP32 GPIO pin assignment for the KANDA Phase 2 embodiment layer")
add_table(
    ["Signal", "GPIO", "Direction", "Notes"],
    [
        ["TRIG_F (Front sensor)",    "GPIO 5",  "Output", "Trigger pulse output"],
        ["ECHO_F (Front sensor)",    "GPIO 34", "Input",  "Input-only pin; 1kΩ+1kΩ divider"],
        ["TRIG_L (Left sensor)",     "GPIO 13", "Output", "Trigger pulse output"],
        ["ECHO_L (Left sensor)",     "GPIO 35", "Input",  "Input-only pin; 1kΩ+1kΩ divider"],
        ["TRIG_R (Right sensor)",    "GPIO 4",  "Output", "Trigger pulse output"],
        ["ECHO_R (Right sensor)",    "GPIO 32", "Input",  "1kΩ+1kΩ divider"],
        ["AIN1 (Motor A dir)",       "GPIO 18", "Output", "TB6612FNG direction bit"],
        ["AIN2 (Motor A dir)",       "GPIO 19", "Output", "TB6612FNG direction bit"],
        ["PWMA (Motor A speed)",     "GPIO 23", "Output", "ledcAttach PWM channel"],
        ["BIN1 (Motor B dir)",       "GPIO 26", "Output", "TB6612FNG direction bit"],
        ["BIN2 (Motor B dir)",       "GPIO 27", "Output", "TB6612FNG direction bit"],
        ["PWMB (Motor B speed)",     "GPIO 14", "Output", "ledcAttach PWM channel"],
        ["SDA (OLED I2C)",           "GPIO 21", "I/O",    "Wire.begin(21, 22) required"],
        ["SCL (OLED I2C)",           "GPIO 22", "Output", "I2C clock line"],
        ["TX (Pi UART bridge)",      "GPIO 1",  "Output", "Phase 3 – reserved"],
        ["RX (Pi UART bridge)",      "GPIO 3",  "Input",  "Phase 3 – reserved"],
    ],
    col_widths=[2.2, 1.0, 1.1, 2.9]
)

heading("5.2 Avoided GPIO Pins")

table_caption("Table 5.2 — ESP32 GPIO pins avoided and reasons for exclusion")
add_table(
    ["GPIO", "Reason Avoided"],
    [
        ["0",    "Strapping pin — controls boot mode; HIGH = normal boot, LOW = download mode"],
        ["2",    "Strapping pin — must be LOW during flash download"],
        ["12",   "Strapping pin — controls flash voltage (VDD_SDIO); incorrect state damages flash"],
        ["15",   "Strapping pin — controls JTAG interface and UART debug output enable"],
        ["6–11", "Internal SPI flash connection — using these pins causes system crash"],
        ["36",   "GPIO36 (VP) — observed instability during testing; avoided as precaution"],
    ],
    col_widths=[1.0, 6.2]
)

heading("5.3 Power Architecture")

para(
    "The power architecture was designed around three principles: stable regulated "
    "supply for the microcontroller, direct high-current supply for the motors, and "
    "physical protection for the battery. The lithium polymer battery was selected "
    "at 7.4 volts nominal (two-cell configuration) to provide sufficient headroom "
    "above the five-volt ESP32 supply for the buck converter to regulate cleanly, "
    "while also providing the higher voltage that improved motor torque."
)

table_caption("Table 5.3 — KANDA power architecture: stages, components, and voltage levels")
add_table(
    ["Stage", "Component", "Voltage In", "Voltage Out", "Notes"],
    [
        ["1", "LiPo Battery",    "—",     "7.4V",  "2S configuration, 2000mAh min"],
        ["2", "BMS",             "7.4V",  "7.4V",  "Overcharge, overdischarge protection"],
        ["3", "Main Switch",     "7.4V",  "7.4V",  "Physical power cutoff"],
        ["4", "Buck Converter",  "7.4V",  "5V",    "Feeds ESP32 VIN pin"],
        ["5", "TB6612FNG VM",    "7.4V",  "7.4V",  "Direct battery feed for motors"],
        ["6", "ESP32 3V3 out",   "5V",    "3.3V",  "Internal regulator for sensors/OLED"],
    ],
    col_widths=[0.8, 2.2, 1.4, 1.4, 3.0]
)

heading("5.4 Signal Conditioning")

para(
    "Each HC-SR04 ECHO pin produced a five-volt pulse when the ultrasonic echo was "
    "received. The ESP32 GPIO input pins had a maximum rated input voltage of "
    "three-point-six volts. Connecting the ECHO pin directly would have applied a "
    "signal forty percent above the rated maximum, risking damage to the ESP32 GPIO "
    "circuitry. A voltage divider was therefore required on each of the three ECHO "
    "lines."
)

para(
    "The voltage divider was implemented using two resistors in series from the "
    "ECHO pin to ground, with the ESP32 GPIO pin connected at the midpoint. "
    "With two equal resistors R, the output voltage was ECHO_voltage times R "
    "divided by 2R, producing exactly half the input voltage. With a five-volt "
    "input, the output was two-point-five volts. The ideal resistor value of two "
    "kiloohms was not available at the time of assembly, and two one-kiloohm "
    "resistors in series were used instead, producing an identical voltage division "
    "ratio. This substitution demonstrated the engineering principle that functional "
    "equivalence is achievable through component substitution when the electrical "
    "behaviour is preserved."
)

heading("5.5 Firmware Architecture")

para(
    "The ESP32 firmware was structured as a single-file Arduino sketch with clearly "
    "separated sections for configuration constants, peripheral function definitions, "
    "and the main loop. All tunable parameters were defined as preprocessor constants "
    "at the top of the file, enabling easy adjustment of obstacle thresholds and "
    "speed values without modifying the control logic."
)

para(
    "The readDistance function used the pulseIn function with a thirty-millisecond "
    "timeout to measure the duration of the ECHO pulse. The timeout prevented the "
    "sensor reading from blocking the main loop when no echo was received, which "
    "could occur if the surface in front of the sensor was angled such that the "
    "ultrasonic pulse was reflected away from the receiver. A return value of "
    "negative one indicated a measurement timeout, which the decision logic treated "
    "as an absence of obstacle at that position."
)

para(
    "The motor control functions, forward, backward, left, right, slightLeft, "
    "slightRight, and stopMotors, were each implemented as short functions that "
    "set the AIN and BIN direction pins and called setSpeed to write the PWM values. "
    "The physical direction of the motors had been reversed in code by swapping the "
    "AIN and AIN2 logic values, correcting for the fact that both motors were mounted "
    "facing inward on the robot chassis, which caused them to spin in opposite "
    "absolute directions for the same pin logic."
)

heading("5.6 LLM Command Protocol")

para(
    "The JSON command protocol between the Raspberry Pi and the ESP32 was designed "
    "to be minimal, unambiguous, and verifiable. Each command was a JSON object "
    "with two fields: action and speed. The action field accepted one of seven "
    "defined string values corresponding to the available movement functions. "
    "The speed field accepted an integer between zero and two hundred and "
    "fifty-five. Commands that contained unrecognised action values or out-of-range "
    "speed values were rejected by the safety validator."
)

table_caption("Table 5.4 — KANDA LLM JSON command protocol: action values and motor behaviours")
add_table(
    ["Action Value", "Motor Behaviour", "Use Case"],
    [
        ["forward",      "Both motors forward at speed",         "Standard navigation"],
        ["backward",     "Both motors reverse at speed",         "Backing away from obstacle"],
        ["left",         "Left motor reverse, right forward",    "Pivot turn left"],
        ["right",        "Left motor forward, right reverse",    "Pivot turn right"],
        ["slight_left",  "Left motor at 50%, right at 100%",     "Gentle left correction"],
        ["slight_right", "Left motor at 100%, right at 50%",     "Gentle right correction"],
        ["stop",         "Both motors off",                      "Safe state / pause"],
    ],
    col_widths=[1.8, 3.0, 2.5]
)

para(
    "Figure 5.1 presents the complete low-level design diagram showing the signal "
    "paths, component connections, and communication interfaces of the KANDA system."
)

figure(IMG_LLD,
       "Figure 5.1 — KANDA Low-Level Design: Pin Mapping, Signal Conditioning, and Interfaces",
       width=Inches(5.8))

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 6
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("CHAPTER 6")
center("PHASE 2 IMPLEMENTATION AND TESTING OF KANDA", size=18)

heading("6.1 Phase 2 Hardware Implementation")

para(
    "The Phase 2 implementation proceeded through the incremental sequence defined "
    "in the methodology. The power architecture was assembled and verified first. "
    "The buck converter output was measured with a multimeter to confirm a stable "
    "five-volt output before the ESP32 was connected. The battery management system "
    "was tested for correct cutoff behaviour by monitoring the output voltage as "
    "the battery was discharged to the cutoff threshold. After confirming stable "
    "power delivery, a minimal blink sketch was flashed to the ESP32 to verify "
    "USB programming connectivity and basic operation."
)

para(
    "The three ultrasonic sensors were then connected, one at a time, with voltage "
    "dividers on each ECHO pin. Each sensor was tested individually using a short "
    "sketch that printed distance readings to the serial monitor at one-second "
    "intervals. Readings were verified to be accurate within approximately one "
    "centimetre for distances between five and one hundred and fifty centimetres, "
    "with degraded accuracy beyond that range. The voltage at each ECHO-to-GPIO "
    "connection point was measured with a multimeter during active sensing to confirm "
    "that the divider was producing the expected two-point-five-volt signal."
)

figure(IMG_ROBOT2,
       "Figure 6.1 — KANDA Phase 2: Front view showing HC-SR04 sensors, differential drive wheels, and assembled chassis",
       width=Inches(5.0))

para(
    "The SSD1306 OLED display was connected to GPIO 21 (SDA) and GPIO 22 (SCL) and "
    "tested with a static text sketch. An important discovery during this step was "
    "that the ESP32 Arduino framework did not default the I2C bus to pins 21 and 22 "
    "without an explicit Wire.begin(21, 22) call in the setup function. Without this "
    "call, the display failed to initialise even though the wiring was correct. "
    "After adding the explicit initialisation call, the display functioned reliably "
    "with the Adafruit SSD1306 library at I2C address 0x3C."
)

figure(IMG_ROBOT1,
       "Figure 6.2 — KANDA Phase 2: Top view showing ESP32 DevKit, TB6612FNG motor driver, OLED display, and battery pack",
       width=Inches(4.5))

para(
    "The TB6612FNG motor driver was connected and the motor control functions were "
    "tested in isolation before integration with the sensor loop. Each of the seven "
    "movement functions was tested by observing wheel rotation direction and speed "
    "for ten seconds, comparing the result to the expected behaviour from the "
    "TB6612FNG truth table. Forward movement was initially observed to move the "
    "robot in the reverse direction, which was corrected by swapping the AIN1 and "
    "AIN2 logic values in the forward and backward functions, reflecting the "
    "inward-facing motor mounting on the chassis."
)

heading("6.2 Key Engineering Decisions and Problem Resolutions")

para(
    "The most significant hardware challenge during Phase 2 was the identification "
    "of the correct voltage divider component values. The ideal two-kiloohm series "
    "resistor was not available in the component inventory at the time of assembly. "
    "Rather than delaying the build, two one-kiloohm resistors were connected in "
    "series to form a two-kiloohm equivalent. The resulting voltage division was "
    "mathematically identical to a single two-kiloohm resistor because the ratio "
    "of the series combination to the lower half of the divider was the same. "
    "This decision was validated by measuring the output voltage and confirmed to "
    "be safe for the ESP32 GPIO inputs."
)

para(
    "The second significant challenge was encountered during firmware compilation "
    "after upgrading to ESP32 Arduino core version 3. The existing firmware used "
    "the ledcSetup and ledcAttachPin functions for PWM configuration, which had "
    "been removed in the version 3 API. The compiler produced errors on these calls "
    "that initially appeared to be library or IDE issues. After reviewing the "
    "ESP32 Arduino core version 3 migration documentation, it was found that both "
    "functions had been replaced by a single ledcAttach function that accepted the "
    "pin number, frequency, and resolution as arguments. The firmware was updated "
    "accordingly, and the PWM output was confirmed to be correct through oscilloscope "
    "measurement of the PWMA and PWMB pins."
)

para(
    "The third challenge was the identification and resolution of GPIO pin conflicts. "
    "An initial pin assignment used GPIO 36 for one of the ECHO signals. During "
    "testing, the readings from this pin were observed to be unstable, producing "
    "intermittent false distance measurements that did not correspond to the actual "
    "sensor output. GPIO 36 was a known problematic pin on some ESP32 DevKit boards "
    "because of its connection to the internal ADC1 circuit. The ECHO signal was "
    "reassigned to GPIO 32, which produced stable and consistent readings."
)

heading("6.3 Obstacle Avoidance Behaviour Validation")

para(
    "The complete obstacle avoidance loop was tested in a corridor environment "
    "with walls on both sides and obstacles placed at various positions and distances. "
    "The robot demonstrated reliable front obstacle detection at the twenty-centimetre "
    "threshold, stopping promptly and turning away from the obstacle based on the "
    "relative left and right distances. The side correction logic, which applied "
    "differential speed reduction when a wall was within fifteen centimetres on "
    "either side, produced smooth steering corrections that prevented the robot "
    "from scraping along walls."
)

para(
    "The OLED display was confirmed to update correctly in real time, showing the "
    "current front, left, and right distances and the active movement decision label "
    "matching the observed robot behaviour. Serial monitor output was cross-referenced "
    "with the OLED display to confirm that both output channels showed identical "
    "sensor readings and decisions. The loop was confirmed to run at the specified "
    "ten-hertz rate without blocking."
)

figure(IMG_FLOW,
       "Figure 6.3 — KANDA sense-decide-act flow: Phase 2 rule-based loop (green) and Phase 3 LLM-driven loop (grey)",
       width=Inches(5.5))

heading("6.4 Phase 3 Design Status")

para(
    "At the time of submission, the Phase 3 implementation had not been started, "
    "but the complete architectural design, software specification, and hardware "
    "component list were finalised and documented. The ESP32 firmware had been "
    "prepared for Phase 3 integration by reserving GPIO pins 1 and 3 for the "
    "Raspberry Pi UART bridge and leaving them unassigned in the Phase 2 pin map. "
    "The JSON command parsing logic was designed and documented in the protocol "
    "specification, ready to be implemented in the ESP32 firmware as an additional "
    "mode activated when a Raspberry Pi connection was detected."
)

para(
    "The Raspberry Pi Python AI pipeline was fully specified in the software "
    "requirement specification and low-level design chapters. The sequence of "
    "initialising the Whisper model, opening the camera stream, setting up the "
    "UART connection, and entering the main context-build-query-validate-command "
    "loop was defined with sufficient detail that implementation could begin "
    "immediately after Phase 2 testing was complete. The hardware description "
    "prompt template was drafted and included in the project documentation."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 7
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("CHAPTER 7")
center("CONCLUSION AND FUTURE WORK", size=18)

heading("7.1 Conclusion")

para(
    "The KANDA project successfully completed Phase 2 of its two-phase development "
    "plan, delivering a fully functional, stable, and validated embodiment layer "
    "for an LLM-powered household companion robot. The ESP32-based hardware platform "
    "demonstrated autonomous obstacle avoidance, smooth differential steering, "
    "real-time OLED feedback, and robust sensor readings in an indoor environment. "
    "The power architecture was designed and validated to provide stable regulated "
    "power to the microcontroller while supplying high-current motor power directly "
    "from the battery, preventing voltage instability during motor operation."
)

para(
    "The project demonstrated that a hardware-first, incremental development "
    "methodology was effective for building complex robotic systems from commodity "
    "components. Each hardware subsystem was verified independently before integration, "
    "maintaining a clear chain of responsibility for any issue that arose and "
    "ensuring that the final integrated system was built on a foundation of "
    "individually validated components. This approach also produced better-documented "
    "hardware designs, since each subsystem was described in terms of its individual "
    "verification results before being combined with others."
)

para(
    "The architectural separation between the ESP32 execution body and the Raspberry "
    "Pi reasoning brain was demonstrated to be a sound engineering choice. It provided "
    "hardware safety through the fallback to rule-based operation during any AI "
    "layer failure, enabled independent development and testing of both layers, and "
    "created a clean interface boundary through which the intelligence layer could "
    "be upgraded without modifying the embodiment layer firmware. The JSON command "
    "protocol with safety validation provided a structured, verifiable interface "
    "that constrained all LLM output to hardware-safe values before execution."
)

para(
    "The revision of the application domain from greenhouse monitoring to household "
    "companion robotics was a critical decision that strengthened the project's "
    "academic contribution. Greenhouse monitoring tasks were largely solvable by "
    "rule-based sensor logic, providing weak justification for LLM integration. "
    "Household companion tasks, including elderly assistance, child tutoring, "
    "medication reminders, home automation, and emergency alerting, required the "
    "contextual reasoning, natural language adaptability, and personalisation that "
    "only an LLM-based system could provide. This domain change aligned the "
    "technical architecture with a compelling social need and produced a stronger, "
    "more defensible research contribution."
)

heading("7.2 Future Work")

para(
    "The immediate priority for future work was the implementation of Phase 3. This "
    "required the physical mounting of the Raspberry Pi on the robot chassis, the "
    "UART wiring connection to GPIO 1 and 3 of the ESP32, and the installation of "
    "the Python AI pipeline on the Pi. The multimodal context builder, LLM API "
    "client, safety validator, and text-to-speech module were all fully specified "
    "and ready to be coded. The first Phase 3 milestone was to establish a reliable "
    "UART data link between the Pi and the ESP32 and verify that JSON commands "
    "transmitted from the Pi were correctly parsed and executed by the ESP32 firmware."
)

para(
    "The second area of future work was the evaluation of the companion use cases "
    "in household settings. A structured evaluation protocol would be designed "
    "covering at minimum three scenarios: an elderly assistance interaction, a "
    "child tutoring session, and a medication reminder delivery. Each scenario "
    "would be evaluated for response latency, command safety validation accuracy, "
    "speech recognition quality, and user interaction naturalness. This evaluation "
    "would provide the evidence base for any publication arising from the project."
)

para(
    "Longer-term future work included the deployment of a quantised local LLM on "
    "the Raspberry Pi to remove the internet connectivity requirement, the addition "
    "of multi-turn conversation memory for personalised interactions, the integration "
    "with smart home APIs for extended home automation capability, and the development "
    "of a semantic safety validation layer that evaluated the contextual "
    "appropriateness of commands in addition to their numerical validity. A user "
    "study with elderly participants was also planned to validate the companion "
    "use cases against real-world requirements."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═════════════════════════════════════════════════════════════════════════════

chapter_label("REFERENCES")
doc.add_paragraph()

refs = [
    "[1]  M. Ahn, A. Brohan, N. Brown, J. Chiang, O. Cortes et al., \"Do As I Can, Not As I "
    "Say: Grounding Language in Robotic Affordances,\" in Proc. 6th Conf. on Robot Learning "
    "(CoRL), Auckland, New Zealand, 2022, pp. 287-318.",

    "[2]  A. Brohan, N. Brown, J. Carbajal, Y. Chebotar, J. Dabis et al., \"RT-2: "
    "Vision-Language-Action Models Transfer Web Knowledge to Robotic Control,\" in Proc. 7th "
    "Conf. on Robot Learning (CoRL), Atlanta, GA, 2023, pp. 2165-2183.",

    "[3]  D. Driess, F. Xia, M. S. M. Sajjadi, C. Lynch, A. Chowdhery et al., \"PaLM-E: "
    "An Embodied Multimodal Language Model,\" in Proc. 40th Int. Conf. on Machine Learning "
    "(ICML), Honolulu, HI, 2023, pp. 8469-8488.",

    "[4]  Z. Yang, Y. Liu, X. Chen and R. Shu, \"RoboGuard: Safety Guardrails for "
    "LLM-Enabled Robotic Systems,\" arXiv:2503.07885 [cs.RO], Mar. 2025.",

    "[5]  X. Liu, Z. Wang and H. Chen, \"SafeEmbodAI: Safety Validation Mechanisms for "
    "Mobile Robot Navigation with Large Language Models,\" arXiv:2409.01630 [cs.RO], "
    "Sep. 2024.",

    "[6]  J. Shen, Y. Zhang, L. Chen and W. Huang, \"Acceptability and Usability of a "
    "Socially Assistive Robot Integrated With a Large Language Model for Enhanced "
    "Human-Robot Interaction in a Geriatric Care Institution: Mixed Methods Evaluation,\" "
    "JMIR Human Factors, vol. 12, no. 1, p. e58553, 2025.",

    "[7]  M. Pinto, A. Costa and R. Sousa, \"Integrating a Large Language Model Into a "
    "Socially Assistive Robot in a Hospital Geriatric Unit: Two-Wave Comparative Study,\" "
    "JMIR Human Factors, vol. 12, no. 2, p. e61207, 2025.",

    "[8]  L. Fischer, T. Weiss and K. Schulz, \"Recommendations for Designing Conversational "
    "Companion Robots With Older Adults Through Foundation Models,\" Frontiers in Robotics "
    "and AI, vol. 11, p. 1342876, 2024.",

    "[9]  Espressif Systems, \"ESP-SparkBot: Large Language Model Robot Powered by ESP32-S3,\" "
    "Espressif Systems Technical Report, Shanghai, China, 2025.",

    "[10] A. Radford, J. W. Kim, T. Xu, G. Brockman, C. McLeavey and I. Sutskever, "
    "\"Robust Speech Recognition via Large-Scale Weak Supervision,\" in Proc. 40th Int. "
    "Conf. on Machine Learning (ICML), Honolulu, HI, 2023, pp. 28492-28518.",

    "[11] Espressif Systems, \"ESP32 Technical Reference Manual,\" Rev. 5.2, Espressif "
    "Systems, Shanghai, China, 2024.",

    "[12] Toshiba Semiconductor, \"TB6612FNG Dual DC Motor Driver IC Datasheet,\" Rev. 2.0, "
    "Toshiba Electronic Devices and Storage Corporation, Tokyo, Japan, 2020.",

    "[13] Solomon Systech, \"SSD1306 OLED Dot Matrix Driver with Controller Datasheet,\" "
    "Rev. 1.1, Solomon Systech Ltd., Hong Kong, 2008.",

    "[14] J. S. Albus, \"Outline for a Theory of Intelligence,\" IEEE Trans. Syst., Man, "
    "Cybern., vol. 21, no. 3, pp. 473-509, May/Jun. 1991.",

    "[15] R. A. Brooks, \"Intelligence Without Representation,\" Artif. Intell., vol. 47, "
    "no. 1-3, pp. 139-159, Jan. 1991.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(ref)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

doc.save("kanda_report.docx")
print("✅  kanda_report.docx saved")
